"""End-to-end link test for the 'scheduled Word delivery' scenario.

Chain under test (cron payload → delivery routing is already covered by
tests/test_scheduler_delivery.py; this file proves the *other* half):

    real .docx generation
      → send_file tool
      → MessageBus.publish_outbound
      → real WeixinChannel.send()  (routes IMAGE/FILE/TEXT blocks)
      → _send_file  (getuploadurl → CDN upload → sendmessage)

Only the iLink HTTP boundary is mocked; docx generation, bus routing, channel
block routing, and message assembly are exercised for real.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from echo_agent.agent.tools.send_file import SendFileTool
from echo_agent.bus.queue import MessageBus
from echo_agent.channels import weixin as wx
from echo_agent.channels.weixin import WeixinChannel
from echo_agent.config.schema import WeixinChannelConfig


def _make_docx(path: Path) -> None:
    """Generate a real .docx the way a cron-triggered agent step would."""
    from docx import Document

    doc = Document()
    doc.add_heading("每日报告", level=1)
    doc.add_paragraph("这是定时任务生成的 Word 文档。")
    doc.save(str(path))


@pytest.mark.asyncio
async def test_scheduled_word_delivery_end_to_end(tmp_path, monkeypatch):
    # ── arrange: real bus + real weixin channel, mock only the iLink HTTP API ──
    captured: dict = {"sendmessage": []}

    async def fake_get_upload_url(session, **kwargs):
        captured["upload_url"] = kwargs
        return {"upload_full_url": "https://novac2c.cdn.weixin.qq.com/c2c/upload"}

    async def fake_upload_ciphertext(session, *, ciphertext, upload_url):
        captured["ciphertext_size"] = len(ciphertext)
        return "ENCRYPTED_PARAM"

    async def fake_api_post(session, *, base_url, endpoint, payload, token, timeout_ms):
        captured["sendmessage"].append(payload)
        return {"errcode": 0}

    monkeypatch.setattr(wx, "_get_upload_url", fake_get_upload_url)
    monkeypatch.setattr(wx, "_upload_ciphertext", fake_upload_ciphertext)
    monkeypatch.setattr(wx, "_api_post", fake_api_post)

    bus = MessageBus()
    cfg = WeixinChannelConfig(
        account_id="acct@im.bot", token="acct@im.bot:tok", data_dir=str(tmp_path / "wx"),
    )
    channel = WeixinChannel(cfg, bus)
    channel._send_session = object()  # truthy; real network is monkeypatched away
    bus.subscribe_outbound(channel.name, channel.send)

    # ── act: agent generates the Word doc, then calls send_file ────────────────
    ws = tmp_path / "ws"
    ws.mkdir()
    docx_path = ws / "daily_report.docx"
    _make_docx(docx_path)

    tool = SendFileTool(str(ws), restrict=True, publish_fn=bus.publish_outbound)
    result = await tool.execute({
        "channel": "weixin",
        "chat_id": "wxid_123",
        "file_path": str(docx_path),
        "caption": "您的每日报告",
    })

    # ── assert: the .docx reached weixin.send() as a file attachment ───────────
    assert result.success, result.error
    assert captured["upload_url"]["media_type"] == wx._MEDIA_FILE  # docx → attachment
    payloads = captured["sendmessage"]
    # a text caption + a file item are both delivered
    file_msgs = [p for p in payloads if p["msg"]["item_list"][0]["type"] == wx._ITEM_FILE]
    assert len(file_msgs) == 1
    file_item = file_msgs[0]["msg"]["item_list"][0]["file_item"]
    assert file_item["file_name"] == "daily_report.docx"
    assert file_msgs[0]["msg"]["to_user_id"] == "wxid_123"
    text_msgs = [p for p in payloads if p["msg"]["item_list"][0]["type"] == wx._ITEM_TEXT]
    assert any("每日报告" in p["msg"]["item_list"][0]["text_item"]["text"] for p in text_msgs)
